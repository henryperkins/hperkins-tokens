import importlib.util
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Pt


SCRIPT_PATH = Path(__file__).resolve().parents[1] / "update-support-resume.py"
SPEC = importlib.util.spec_from_file_location("update_support_resume", SCRIPT_PATH)
UPDATER = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(UPDATER)


def add_sized_hyperlink(paragraph, text, url, size, size_tag="w:sz"):
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    size_element = OxmlElement(size_tag)
    size_element.set(qn("w:val"), str(int(size * 2)))
    run_properties.append(size_element)
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


class MinimumEffectiveBodySizeTests(unittest.TestCase):
    def make_document(self):
        document = Document()
        document.styles["Normal"].font.size = Pt(9.5)
        return document

    def test_detects_undersized_direct_run(self):
        document = self.make_document()
        paragraph = document.add_paragraph()
        paragraph.add_run("too small").font.size = Pt(8)
        self.assertEqual(UPDATER.minimum_effective_body_size(document), 8.0)
        with self.assertRaisesRegex(ValueError, r"8\.0pt$"):
            UPDATER.assert_minimum_body_size(document)

    def test_detects_undersized_hyperlink_run(self):
        document = self.make_document()
        paragraph = document.add_paragraph()
        add_sized_hyperlink(paragraph, "small link", "https://example.test/", 8.5, "w:szCs")
        self.assertEqual(UPDATER.minimum_effective_body_size(document), 8.5)
        with self.assertRaisesRegex(ValueError, r"8\.5pt$"):
            UPDATER.assert_minimum_body_size(document)

    def test_requires_exactly_one_scoped_resume_event(self):
        document = self.make_document()
        event_style = document.styles.add_style("Resume Event", 1)
        event_style.font.size = Pt(8.5)
        document.add_paragraph(UPDATER.EVENT, style=event_style)
        self.assertEqual(UPDATER.assert_resume_event_contract(document), 8.5)

        document.add_paragraph("unexpected second event", style=event_style)
        with self.assertRaisesRegex(ValueError, r"exactly one Resume Event"):
            UPDATER.assert_resume_event_contract(document)

    def test_requires_exact_event_copy_and_minimum_size(self):
        wrong_copy = self.make_document()
        wrong_style = wrong_copy.styles.add_style("Resume Event", 1)
        wrong_style.font.size = Pt(8.5)
        wrong_copy.add_paragraph("unexpected event", style=wrong_style)
        with self.assertRaisesRegex(ValueError, r"approved WCUS event copy$"):
            UPDATER.assert_resume_event_contract(wrong_copy)

        undersized = self.make_document()
        undersized_style = undersized.styles.add_style("Resume Event", 1)
        undersized_style.font.size = Pt(8)
        undersized.add_paragraph(UPDATER.EVENT, style=undersized_style)
        with self.assertRaisesRegex(ValueError, r"below 8\.5pt: 8\.0pt$"):
            UPDATER.assert_resume_event_contract(undersized)


if __name__ == "__main__":
    unittest.main()
