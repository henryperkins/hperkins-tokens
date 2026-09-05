#!/usr/bin/env python3
"""Surgically refresh the existing one-page support résumé DOCX."""

from __future__ import annotations

import datetime
import io
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "assets" / "documents" / "henry-perkins-wordpress-support-engineer-resume.docx"
TITLE = "Henry Perkins — WordPress Support Engineer"
AUTHOR = "Henry Perkins"
# Document metadata only; the revision date never appears in the visible text.
REVISED = datetime.datetime(2026, 9, 4, tzinfo=datetime.timezone.utc)
REVISED_LABEL = "September 4, 2026"
# Post-event copy, mirroring the published About page. The pre-event
# "Selected to staff" line is retired and forbidden by the artifact contracts,
# so a binary that was not rebuilt cannot pass them.
EVENT = (
    "WORDCAMP US 2026 — Phoenix · Staffed the Core AI booth, walking maintainers "
    "and agency developers through AI provider tooling"
)
HEADLINE = "WordPress · Gutenberg · REST/HTTP/DNS · Defect reproduction · Fix validation · Customer communication"
SUMMARY = (
    "WordPress support professional with prior WordPress.com support experience, current independent "
    "client delivery, and upstream WordPress contributions. Reproduces defects, tests fixes, and reports "
    "findings to customers and engineering."
)

# Type scale in points. Body text is 10.5pt; the event line is the one
# deliberately smaller run and carries its own floor in assert_resume_event_contract.
BODY_SIZE = 10.5
TITLE_SIZE = 17.5
SECTION_SIZE = 10.5
EVENT_SIZE = 9.0

# Every entry is (paragraph style, [(visible text, hyperlink URL or None), ...]).
# A Resume Body whose first plain segment ends in " — " renders that segment as a
# bold lead-in label; keep labels in their own segment so the body stays regular.
RESUME = [
    ("Heading 1", [(TITLE, None)]),
    ("Normal", [(HEADLINE, None)]),
    ("Normal", [
        ("Chicago, IL  ·  ", None),
        ("htperkins@gmail.com", "mailto:htperkins@gmail.com"),
        ("  ·  ", None),
        ("hperkins.blog", "https://hperkins.blog/"),
        ("  ·  ", None),
        ("GitHub", "https://github.com/henryperkins"),
    ]),
    ("Resume Event", [(EVENT, None)]),
    ("Normal", [(SUMMARY, None)]),
    ("Heading 2", [("EXPERIENCE", None)]),
    ("Resume Entry", [("Independent Technology Consultant — Lakefront Digital  |  Oct 2022–Present · Chicago, remote", None)]),
    ("Resume Body", [
        ("Delivery scope — ", None),
        ("WordPress builds, API integrations, and documentation from discovery through post-launch support.", None),
    ]),
    ("Resume Body", [
        ("Client launch — ", None),
        ("Delivered ", None),
        ("DJ Lee & Voices of Judah", "https://thevoicesofjudah.com/"),
        ("’s booking-focused website from discovery through launch, including server-side booking validation and ", None),
        ("publicly available source code", "https://github.com/henryperkins/dj-judas-v2"),
        (".", None),
    ]),
    ("Resume Entry", [("Automattic — Happiness Engineer  |  WordPress.com · Remote · Oct–Nov 2012", None)]),
    ("Resume Body", [(
        "Resolved WordPress.com issues across publishing, configuration, billing, domains, and DNS, turning recurring "
        "problems into documentation and reproducible bug reports for customers, product, and engineering.",
        None,
    )]),
    ("Heading 2", [("SELECTED WORDPRESS INVESTIGATIONS & CONTRIBUTIONS", None)]),
    ("Resume Body", [
        ("Defect reproduction — ", None),
        ("Reported and reproduced a WordPress AI Guidelines defect (", None),
        ("Issue #529", "https://github.com/WordPress/ai/issues/529"),
        ("): an artifact guideline could shadow the content guideline. A maintainer’s fix, ", None),
        ("PR #593", "https://github.com/WordPress/ai/pull/593"),
        (", shipped in ", None),
        ("WordPress AI 1.0.1", "https://github.com/WordPress/ai/releases/tag/1.0.1"),
        (".", None),
    ]),
    ("Resume Body", [
        ("Fix validation — ", None),
        ("Reported a WordPress AI request-logging gap (", None),
        ("Issue #732", "https://github.com/WordPress/ai/issues/732"),
        ("); integration-tested Anubhav Anand’s proposed fix (", None),
        ("PR #757", "https://github.com/WordPress/ai/pull/757"),
        (", open), found duplicate successes and missing failures, and proposed the ownership split.", None),
    ]),
    ("Resume Body", [
        ("Documentation — ", None),
        ("Wrote and refined the Content Resizing and Title Generation experiment documentation in ", None),
        ("WordPress/ai PR #501", "https://github.com/WordPress/ai/pull/501"),
        ("; merged May 18, 2026 and credited in the 1.0.0 release notes.", None),
    ]),
    ("Resume Body", [
        ("Compatibility fix — ", None),
        ("Directed and reviewed an AI-assisted fix so the OpenAI provider advertises sampling options only for models that accept them, with tests: ", None),
        ("WordPress/ai-provider-for-openai PR #40", "https://github.com/WordPress/ai-provider-for-openai/pull/40"),
        (", merged Aug 16, 2026.", None),
    ]),
    ("Resume Body", [
        ("Input validation — ", None),
        ("Directed and reviewed an AI-assisted contribution that rejects NAN and infinite embedding values, with a regression test for each: ", None),
        ("WordPress/php-ai-client PR #263", "https://github.com/WordPress/php-ai-client/pull/263"),
        (", open contribution.", None),
    ]),
    ("Heading 2", [("SELECTED PROJECTS", None)]),
    ("Resume Body", [("Solo projects, built AI-assisted under my direction and review, with public tagged releases.", None)]),
    ("Resume Entry", [("Flavor Agent — Creator  |  WordPress agent-governance plugin", None)]),
    ("Resume Body", [
        ("v0.1.0", "https://github.com/henryperkins/flavor-agent/releases/tag/v0.1.0"),
        (
            " released Aug 26, 2026: puts AI-proposed site changes through validation, admin approval, an audit record, "
            "and safe undo inside the block editor and wp-admin.",
            None,
        ),
    ]),
    ("Resume Entry", [("AI Provider for Codex — Creator  |  independent WordPress plugin", None)]),
    ("Resume Body", [
        ("v2.1", "https://github.com/henryperkins/ai-provider-for-codex/releases/tag/v2.1"),
        (" released: connects Codex text and image generation to the WordPress AI Client through a local sidecar.", None),
    ]),
    ("Resume Entry", [("HPerkins Tokens — Creator  |  WordPress block theme behind hperkins.blog", None)]),
    ("Resume Body", [
        ("v0.3.53", "https://github.com/henryperkins/hperkins-tokens/releases/tag/v0.3.53"),
        (" released · ", None),
        ("hperkins.blog", "https://hperkins.blog/"),
        (
            ": a token-governed block theme in which editors choose only named design tokens, "
            "guarded by verifier scripts for content, typography, and accessibility checks.",
            None,
        ),
    ]),
    ("Heading 2", [("TECHNICAL SKILLS & ADDITIONAL EXPERIENCE", None)]),
    ("Resume Body", [(
        "Code investigation and review: PHP, JavaScript, Gutenberg, REST API, WP-CLI · "
        "Support: HTTP, DNS, CSS, browser debugging, escalation triage · "
        "Tooling: Git/GitHub, GitHub Actions, Plugin Check, PHPStan, Cloudflare Workers.",
        None,
    )]),
    ("Resume Body", [
        ("Developer community — ", None),
        (
            "PageLines Developer Community Manager (May–Oct 2012): onboarding content, tutorials, and day-to-day "
            "developer relations, turning community feedback into clearer product guidance.",
            None,
        ),
    ]),
    ("Resume Body", [
        ("Customer and operations roles — ", None),
        (
            "Starbucks Shift Supervisor (2019–2022); Sodexo Starbucks Manager (2018–2019); "
            "Clinique Consultant (2015–2017); Micro Center Customer Service/Sales (2009–2012).",
            None,
        ),
    ]),
]


def first_run_properties(paragraph):
    for run in paragraph._p.xpath(".//w:r"):
        run_properties = run.find(qn("w:rPr"))
        if run_properties is not None:
            return deepcopy(run_properties)
    return None


def paragraph_by_style(document, names):
    for paragraph in document.paragraphs:
        if paragraph.style.name in names:
            return paragraph
    raise ValueError(f"Résumé has no paragraph using {', '.join(names)}")


def plain_body_run_properties(document):
    """Run properties of plain body text: the first Resume Body paragraph with no lead-in label."""
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Resume Body" and " — " not in paragraph.text:
            return first_run_properties(paragraph)
    return None


def label_body_run_properties(document):
    """Run properties of the bold lead-in label that opens a Resume Body paragraph."""
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Resume Body" and " — " in paragraph.text:
            return first_run_properties(paragraph)
    return None


def copy_paragraph_format(source, destination):
    source_properties = source._p.pPr
    if source_properties is None:
        return
    destination_properties = destination._p.get_or_add_pPr()
    for child in list(destination_properties):
        if child.tag != qn("w:pStyle"):
            destination_properties.remove(child)
    for child in source_properties:
        if child.tag != qn("w:pStyle"):
            destination_properties.append(deepcopy(child))


def apply_run_properties(run, properties):
    if properties is None:
        return
    existing = run._r.rPr
    if existing is not None:
        run._r.remove(existing)
    run._r.insert(0, deepcopy(properties))


def without_literal_sizes(run_properties):
    """Copy run properties minus any literal size, so the run inherits its style's scale."""
    if run_properties is None:
        return None
    properties = deepcopy(run_properties)
    for size_name in ("w:sz", "w:szCs"):
        size_element = properties.find(qn(size_name))
        if size_element is not None:
            properties.remove(size_element)
    return properties


def add_hyperlink(paragraph, text, url, run_properties):
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    if run_properties is not None:
        run.append(deepcopy(run_properties))
    text_node = OxmlElement("w:t")
    if text[:1].isspace() or text[-1:].isspace():
        text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_style_visuals(style, normal_style, *, size, bold, color, paragraph_source=None):
    style.base_style = normal_style
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    if paragraph_source is not None:
        source_format = paragraph_source.paragraph_format
        target_format = style.paragraph_format
        for attribute in (
            "space_before", "space_after", "line_spacing", "keep_with_next",
            "keep_together", "page_break_before", "widow_control",
        ):
            setattr(target_format, attribute, getattr(source_format, attribute))


def remove_external_hyperlink_relationships(document):
    relationships = document.part.rels
    for relationship_id, relationship in list(relationships.items()):
        if relationship.reltype == RELATIONSHIP_TYPE.HYPERLINK and relationship.is_external:
            del relationships[relationship_id]


def effective_style_size(style):
    while style is not None:
        if style.font.size is not None:
            return style.font.size.pt
        style = style.base_style
    return None


def effective_run_size(document, paragraph, run_element):
    run_properties = run_element.find(qn("w:rPr"))
    if run_properties is not None:
        direct_sizes = []
        for size_name in ("w:sz", "w:szCs"):
            size_element = run_properties.find(qn(size_name))
            if size_element is not None:
                direct_sizes.append(int(size_element.get(qn("w:val"))) / 2)
        if direct_sizes:
            return min(direct_sizes)

        run_style = run_properties.find(qn("w:rStyle"))
        if run_style is not None:
            style_id = run_style.get(qn("w:val"))
            for style in document.styles:
                if style.type == WD_STYLE_TYPE.CHARACTER and style.style_id == style_id:
                    style_size = effective_style_size(style)
                    if style_size is not None:
                        return style_size

    paragraph_size = effective_style_size(paragraph.style)
    if paragraph_size is not None:
        return paragraph_size
    return effective_style_size(document.styles["Normal"])


def minimum_effective_body_size(document):
    sizes = []
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Resume Event":
            continue
        for run_element in paragraph._p.xpath(".//w:r"):
            if not any(node.text for node in run_element.xpath(".//w:t")):
                continue
            size = effective_run_size(document, paragraph, run_element)
            if size is None:
                raise ValueError(f"Cannot resolve effective font size in paragraph: {paragraph.text}")
            sizes.append(size)
    if not sizes:
        raise ValueError("Résumé contains no non-event text runs to audit")
    return min(sizes)


def assert_minimum_body_size(document, floor=BODY_SIZE):
    minimum = minimum_effective_body_size(document)
    if minimum < floor:
        raise ValueError(f"Résumé effective non-event text falls below {floor:.1f}pt: {minimum:.1f}pt")
    return minimum


def assert_resume_event_contract(document, floor=8.5):
    event_paragraphs = [
        paragraph for paragraph in document.paragraphs
        if paragraph.style.name == "Resume Event"
    ]
    if len(event_paragraphs) != 1:
        raise ValueError(f"Résumé must contain exactly one Resume Event paragraph; found {len(event_paragraphs)}")

    paragraph = event_paragraphs[0]
    visible_text = "".join(
        node.text or "" for node in paragraph._p.xpath(".//w:t")
    )
    if visible_text != EVENT:
        raise ValueError("Resume Event paragraph does not match the approved WCUS event copy")

    sizes = [
        effective_run_size(document, paragraph, run_element)
        for run_element in paragraph._p.xpath(".//w:r")
        if any(node.text for node in run_element.xpath(".//w:t"))
    ]
    if not sizes or any(size is None for size in sizes):
        raise ValueError("Cannot resolve the Resume Event effective font size")
    minimum = min(sizes)
    if minimum < floor:
        raise ValueError(f"Resume Event text falls below {floor:.1f}pt: {minimum:.1f}pt")
    return minimum


def canonical_docx_bytes(document):
    source = io.BytesIO()
    document.save(source)
    source.seek(0)
    output = io.BytesIO()
    with zipfile.ZipFile(source, "r") as archive, zipfile.ZipFile(
        output, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
    ) as canonical:
        for name in sorted(archive.namelist()):
            info = zipfile.ZipInfo(name, (1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 0
            info.external_attr = 0
            canonical.writestr(info, archive.read(name))
    return output.getvalue()


def main():
    if not DOCX_PATH.is_file():
        raise SystemExit(f"Existing résumé DOCX not found: {DOCX_PATH}")

    document = Document(DOCX_PATH)
    if len(document.sections) != 1:
        raise SystemExit(f"Expected one section, found {len(document.sections)}")

    first_paragraph = document.paragraphs[0]
    section_paragraph = paragraph_by_style(document, {"Resume Section", "Heading 2"})
    entry_paragraph = paragraph_by_style(document, {"Resume Entry"})
    body_paragraph = paragraph_by_style(document, {"Resume Body"})

    title_run_properties = first_run_properties(first_paragraph)
    section_run_properties = first_run_properties(section_paragraph)
    entry_run_properties = first_run_properties(entry_paragraph)
    # Plain body runs and bold lead-in labels are read from paragraphs of their
    # own kind, so the choice does not depend on which Resume Body comes first.
    body_run_properties = plain_body_run_properties(document)
    label_run_properties = label_body_run_properties(document)
    if label_run_properties is None:
        label_run_properties = section_run_properties

    link_run_properties = None
    for hyperlink_run in document.element.body.xpath(".//w:hyperlink/w:r"):
        candidate = hyperlink_run.find(qn("w:rPr"))
        if candidate is not None:
            # Links inherit the body scale from the style; a literal size copied
            # from an older document would silently pin them to that older scale.
            link_run_properties = without_literal_sizes(candidate)
            break

    styles = document.styles
    heading_one = styles["Heading 1"]
    heading_two = styles["Heading 2"]
    normal_style = styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(BODY_SIZE)
    set_style_visuals(heading_one, normal_style, size=TITLE_SIZE, bold=True, color="183F5C", paragraph_source=first_paragraph)
    set_style_visuals(heading_two, normal_style, size=SECTION_SIZE, bold=True, color="183F5C", paragraph_source=section_paragraph)
    # Match the regenerated paragraphs immediately, so an older input does not
    # leave style spacing that changes again on the next run.
    heading_two.paragraph_format.space_before = Pt(6)
    heading_two.paragraph_format.space_after = Pt(2)
    if "Resume Event" in styles:
        event_style = styles["Resume Event"]
    else:
        event_style = styles.add_style("Resume Event", WD_STYLE_TYPE.PARAGRAPH)
    set_style_visuals(event_style, normal_style, size=EVENT_SIZE, bold=True, color="9A7530")
    event_style.paragraph_format.space_before = Pt(0)
    event_style.paragraph_format.space_after = Pt(1.2)
    event_style.paragraph_format.line_spacing = 1.0
    event_style.paragraph_format.keep_together = True

    remove_external_hyperlink_relationships(document)
    for paragraph in list(document.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)

    for style_name, segments in RESUME:
        paragraph = document.add_paragraph(style=style_name)
        paragraph_text = "".join(visible_text for visible_text, _url in segments)
        is_headline = style_name == "Normal" and paragraph_text == HEADLINE
        is_contact = style_name == "Normal" and paragraph_text.startswith("Chicago, IL")
        is_summary = style_name == "Normal" and paragraph_text == SUMMARY
        if style_name == "Heading 1":
            copy_paragraph_format(first_paragraph, paragraph)
        elif style_name == "Heading 2":
            copy_paragraph_format(section_paragraph, paragraph)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(2)
        elif style_name == "Resume Entry":
            copy_paragraph_format(entry_paragraph, paragraph)
            paragraph.paragraph_format.space_before = Pt(3)
            paragraph.paragraph_format.space_after = Pt(0.5)
        elif style_name == "Resume Body":
            copy_paragraph_format(body_paragraph, paragraph)
            paragraph.paragraph_format.space_after = Pt(2)
        elif is_summary:
            paragraph.paragraph_format.space_after = Pt(3)

        for segment_index, (visible_text, url) in enumerate(segments):
            if url is not None:
                add_hyperlink(paragraph, visible_text, url, link_run_properties)
                continue

            if style_name == "Resume Entry" and "  |  " in visible_text:
                primary, metadata = visible_text.split("  |  ", 1)
                primary_run = paragraph.add_run(primary)
                apply_run_properties(primary_run, entry_run_properties)
                metadata_run = paragraph.add_run(f"  |  {metadata}")
                metadata_run.font.color.rgb = RGBColor.from_string("525D6B")
                continue

            run = paragraph.add_run(visible_text)
            if style_name == "Heading 1":
                apply_run_properties(run, title_run_properties)
            elif style_name == "Heading 2":
                apply_run_properties(run, section_run_properties)
            elif style_name == "Resume Event":
                run.font.name = "Calibri"
                run.font.size = Pt(EVENT_SIZE)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string("9A7530")
            elif style_name == "Resume Body":
                is_label = segment_index == 0 and visible_text.endswith(" — ")
                apply_run_properties(run, label_run_properties if is_label else body_run_properties)
            elif is_headline:
                run.font.bold = True
            elif is_contact:
                run.font.color.rgb = RGBColor.from_string("525D6B")

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.60)
    section.right_margin = Inches(0.60)
    document.core_properties.title = TITLE
    document.core_properties.author = AUTHOR
    # python-docx exposes dc:description as `comments`.
    document.core_properties.comments = f"Evidence-bounded résumé revised {REVISED_LABEL}."
    document.core_properties.keywords = (
        "WordPress, support engineer, Gutenberg, REST, HTTP, DNS, defect reproduction, fix validation"
    )
    document.core_properties.modified = REVISED

    assert_resume_event_contract(document)
    assert_minimum_body_size(document)
    output = canonical_docx_bytes(document)
    changed = DOCX_PATH.read_bytes() != output
    if changed:
        temporary = DOCX_PATH.with_suffix(".docx.tmp")
        temporary.write_bytes(output)
        temporary.replace(DOCX_PATH)

    reopened = Document(DOCX_PATH)
    visible_text = " ".join(paragraph.text for paragraph in reopened.paragraphs)
    assert_resume_event_contract(reopened)
    minimum_body_size = assert_minimum_body_size(reopened)
    print(f"updated={str(changed).lower()}")
    print(f"sections={len(reopened.sections)}")
    print("page=US Letter 8.50x11.00in")
    margins = reopened.sections[0]
    print(
        "margins="
        f"top:{margins.top_margin.inches:.2f}in bottom:{margins.bottom_margin.inches:.2f}in "
        f"left:{margins.left_margin.inches:.2f}in right:{margins.right_margin.inches:.2f}in"
    )
    print(f"minimum_body_size={minimum_body_size:.1f}pt")
    print(f"event_count={visible_text.count(EVENT)}")


if __name__ == "__main__":
    main()
